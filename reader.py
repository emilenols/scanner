"""
reader.py — format-aware document reader shared by all scanners.

Gemini reads PDF and images natively; it does NOT accept Word/Excel binaries.
So we send PDF/images as bytes, and extract text locally from Word/Excel.
Genuinely unreadable files (legacy .doc, CAD, unsupported images) fail FAST
with a logged reason — never the slow retry loop.

The read-only service account cannot convert files in Drive, so there is no
PDF-conversion path; this is the correct, no-write approach.
"""
import io

# Gemini-native media
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
SUPPORTED_IMAGE_MIMES = {
    "image/png", "image/jpeg", "image/jpg", "image/webp", "image/heic", "image/heif",
}
MAX_TEXT_CHARS = 200_000
MAX_XLSX_ROWS = 200
# Files larger than this are skipped in the synchronous path (discovery/pilot):
# downloading + inline-sending multi-MB scans is slow and times out. They are
# NOT lost — the async Batch (full Pass 1) still classifies them.
MAX_INLINE_BYTES = 10 * 1024 * 1024  # 10 MB


class UnreadableDocument(Exception):
    def __init__(self, reason):
        self.reason = reason
        super().__init__(reason)


def is_eligible(record, include_images=True) -> bool:
    """Can this file be read for classification? (pure — no heavy imports)"""
    if record.get("skip_reason"):
        return False
    fb, mime = record["format_bucket"], record["mime_type"]
    if fb == "PDF":
        return True
    if fb == "Word":
        return mime == DOCX_MIME            # legacy .doc not supported
    if fb == "Excel":
        return mime == XLSX_MIME            # legacy .xls not supported
    if fb == "Text":
        return True
    if fb == "Image":
        return include_images and mime in SUPPORTED_IMAGE_MIMES
    return False


def ineligible_reason(record, include_images=True) -> str:
    fb, mime = record["format_bucket"], record["mime_type"]
    if fb == "Image" and not include_images:
        return "images_excluded_by_config"
    if fb == "Image":
        return "image_format_unsupported"
    if fb == "Word":
        return "legacy_doc_unsupported"
    if fb == "Excel":
        return "legacy_xls_unsupported"
    return f"unsupported_format_{fb.lower()}"


def _download(drive, file_id) -> bytes:
    from googleapiclient.http import MediaIoBaseDownload
    req = drive.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    dl = MediaIoBaseDownload(buf, req)
    done = False
    while not done:
        _, done = dl.next_chunk()
    return buf.getvalue()


def _docx_text(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _xlsx_text(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# Sheet: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= MAX_XLSX_ROWS:
                out.append("... (truncated)")
                break
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def read_for_model(drive, record, include_images=True, enforce_size_cap=True) -> dict:
    """Return {'kind':'media','mime':..,'bytes':..} or {'kind':'text','text':..}.
    Raise UnreadableDocument(reason) for skip-and-log (fast, no retry).

    enforce_size_cap: in the sync path (discovery/pilot) skip files over
    MAX_INLINE_BYTES. The full Batch pass sets this False (it must read all)."""
    if not is_eligible(record, include_images):
        raise UnreadableDocument(ineligible_reason(record, include_images))
    if enforce_size_cap and record.get("size_bytes", 0) > MAX_INLINE_BYTES:
        raise UnreadableDocument("too_large_for_inline")
    fb, mime = record["format_bucket"], record["mime_type"]
    try:
        data = _download(drive, record["drive_id"])
    except UnreadableDocument:
        raise
    except Exception:
        raise UnreadableDocument("download_failed")

    if fb == "PDF":
        return {"kind": "media", "mime": "application/pdf", "bytes": data}
    if fb == "Image":
        return {"kind": "media", "mime": mime, "bytes": data}
    if fb == "Word":
        try:
            txt = _docx_text(data)
        except Exception:
            raise UnreadableDocument("docx_parse_error")
        if len(txt.strip()) < 20:
            raise UnreadableDocument("word_no_extractable_text")  # image-only doc
        return {"kind": "text", "text": txt[:MAX_TEXT_CHARS]}
    if fb == "Excel":
        try:
            txt = _xlsx_text(data)
        except Exception:
            raise UnreadableDocument("xlsx_parse_error")
        if len(txt.strip()) < 5:
            raise UnreadableDocument("xlsx_empty")
        return {"kind": "text", "text": txt[:MAX_TEXT_CHARS]}
    if fb == "Text":
        return {"kind": "text", "text": data.decode("utf-8", "ignore")[:MAX_TEXT_CHARS]}
    raise UnreadableDocument("unsupported")


def sync_contents(result, prompt):
    """Build the genai `contents` list for a synchronous call."""
    if result["kind"] == "media":
        from google.genai import types
        return [types.Part.from_bytes(data=result["bytes"], mime_type=result["mime"]), prompt]
    return [f"Document content:\n\n{result['text']}", prompt]


def batch_parts(result, prompt):
    """Build the `parts` list for a Batch API JSONL request."""
    import base64
    if result["kind"] == "media":
        b64 = base64.b64encode(result["bytes"]).decode()
        return [{"inline_data": {"mime_type": result["mime"], "data": b64}},
                {"text": prompt}]
    return [{"text": f"Document content:\n\n{result['text']}\n\n{prompt}"}]
